#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script pour créer une présentation Word du projet CleanCampusManager
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os
from datetime import datetime

def create_clean_campus_presentation():
    """Créer une présentation Word complète du projet CleanCampusManager"""
    
    # Créer un nouveau document
    doc = Document()
    
    # Configuration des styles
    setup_document_styles(doc)
    
    # Page de titre
    create_title_page(doc)
    
    # Table des matières
    create_table_of_contents(doc)
    
    # Introduction
    create_introduction_section(doc)
    
    # Présentation du projet
    create_project_overview_section(doc)
    
    # Architecture technique
    create_technical_architecture_section(doc)
    
    # Fonctionnalités principales
    create_main_features_section(doc)
    
    # Interface utilisateur
    create_user_interface_section(doc)
    
    # Modèles de données
    create_data_models_section(doc)
    
    # Services et gestion des données
    create_services_section(doc)
    
    # Système de gamification
    create_gamification_section(doc)
    
    # Installation et utilisation
    create_installation_section(doc)
    
    # Équipe de développement
    create_team_section(doc)
    
    # Conclusion
    create_conclusion_section(doc)
    
    # Sauvegarder le document
    filename = f"CleanCampusManager_Presentation_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(filename)
    print(f"Présentation créée avec succès: {filename}")
    return filename

def setup_document_styles(doc):
    """Configurer les styles du document"""
    
    # Style pour les titres principaux
    title_style = doc.styles['Title']
    title_style.font.size = Pt(24)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor(0, 51, 102)
    
    # Style pour les titres de section
    heading_style = doc.styles['Heading 1']
    heading_style.font.size = Pt(18)
    heading_style.font.bold = True
    heading_style.font.color.rgb = RGBColor(0, 102, 204)
    
    # Style pour les sous-titres
    subheading_style = doc.styles['Heading 2']
    subheading_style.font.size = Pt(14)
    subheading_style.font.bold = True
    subheading_style.font.color.rgb = RGBColor(51, 153, 255)
    
    # Style pour le texte normal
    normal_style = doc.styles['Normal']
    normal_style.font.size = Pt(11)
    normal_style.font.name = 'Calibri'

def create_title_page(doc):
    """Créer la page de titre"""
    
    # Titre principal
    title = doc.add_heading('CleanCampus Manager', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Sous-titre
    subtitle = doc.add_paragraph('Système de Gestion du Nettoyage des Résidences Universitaires')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(16)
    subtitle.runs[0].font.italic = True
    
    # Espace
    doc.add_paragraph()
    
    # Description
    description = doc.add_paragraph('Application desktop moderne développée en Python avec Tkinter pour automatiser et optimiser la gestion du nettoyage dans les résidences universitaires.')
    description.alignment = WD_ALIGN_PARAGRAPH.CENTER
    description.runs[0].font.size = Pt(12)
    
    # Espace
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Date
    date_para = doc.add_paragraph(f'Présentation créée le {datetime.now().strftime("%d/%m/%Y")}')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.runs[0].font.size = Pt(10)
    date_para.runs[0].font.italic = True
    
    # Saut de page
    doc.add_page_break()

def create_table_of_contents(doc):
    """Créer la table des matières"""
    
    title = doc.add_heading('Table des Matières', 1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Liste des sections
    sections = [
        "1. Introduction",
        "2. Présentation du Projet",
        "3. Architecture Technique",
        "4. Fonctionnalités Principales",
        "5. Interface Utilisateur",
        "6. Modèles de Données",
        "7. Services et Gestion des Données",
        "8. Système de Gamification",
        "9. Installation et Utilisation",
        "10. Équipe de Développement",
        "11. Conclusion"
    ]
    
    for section in sections:
        p = doc.add_paragraph(section)
        p.style = 'List Bullet'
    
    doc.add_page_break()

def create_introduction_section(doc):
    """Créer la section introduction"""
    
    title = doc.add_heading('1. Introduction', 1)
    
    doc.add_paragraph()
    
    # Contexte
    context = doc.add_heading('Contexte', 2)
    context_text = doc.add_paragraph(
        'Les résidences universitaires font face à des défis quotidiens pour maintenir un environnement propre et sain. '
        'La gestion manuelle des tâches de nettoyage, la répartition des responsabilités et le suivi des performances '
        'représentent des défis majeurs pour les administrateurs et les étudiants.'
    )
    
    doc.add_paragraph()
    
    # Problématique
    problem = doc.add_heading('Problématique', 2)
    problem_text = doc.add_paragraph(
        '• Gestion manuelle et chronophage des plannings de nettoyage\n'
        '• Difficulté à assurer une rotation équitable des tâches\n'
        '• Manque de motivation des étudiants pour participer\n'
        '• Absence de suivi des performances et de la qualité\n'
        '• Communication inefficace entre les parties prenantes'
    )
    problem_text.style = 'List Bullet'
    
    doc.add_paragraph()
    
    # Solution proposée
    solution = doc.add_heading('Solution Proposée', 2)
    solution_text = doc.add_paragraph(
        'CleanCampus Manager est une application desktop moderne qui automatise et optimise '
        'la gestion du nettoyage des résidences universitaires. Elle offre une solution complète '
        'pour la planification, le suivi et la motivation des étudiants.'
    )

def create_project_overview_section(doc):
    """Créer la section présentation du projet"""
    
    doc.add_page_break()
    title = doc.add_heading('2. Présentation du Projet', 1)
    
    doc.add_paragraph()
    
    # Vue d'ensemble
    overview = doc.add_heading('Vue d\'Ensemble', 2)
    overview_text = doc.add_paragraph(
        'CleanCampus Manager est une application de gestion complète développée en Python avec Tkinter. '
        'Elle modernise et automatise la gestion du nettoyage dans les résidences universitaires en offrant '
        'des outils avancés pour la planification, le suivi et la motivation.'
    )
    
    doc.add_paragraph()
    
    # Objectifs
    objectives = doc.add_heading('Objectifs du Projet', 2)
    objectives_list = [
        'Automatiser la création et la gestion des plannings de nettoyage',
        'Assurer une répartition équitable des tâches entre les étudiants',
        'Améliorer la motivation grâce à un système de gamification',
        'Faciliter le suivi des performances et de la qualité',
        'Optimiser la communication entre administrateurs, gestionnaires et étudiants',
        'Réduire la charge administrative de gestion des résidences'
    ]
    
    for obj in objectives_list:
        p = doc.add_paragraph(obj)
        p.style = 'List Bullet'
    
    doc.add_paragraph()
    
    # Avantages
    benefits = doc.add_heading('Avantages', 2)
    benefits_list = [
        'Interface utilisateur moderne et intuitive',
        'Automatisation complète des processus',
        'Système de gamification pour motiver les étudiants',
        'Gestion multi-niveaux (Administrateur, Gestionnaire, Étudiant)',
        'Export de données et génération de rapports',
        'Sauvegarde et restauration automatiques des données'
    ]
    
    for benefit in benefits_list:
        p = doc.add_paragraph(benefit)
        p.style = 'List Bullet'

def create_technical_architecture_section(doc):
    """Créer la section architecture technique"""
    
    doc.add_page_break()
    title = doc.add_heading('3. Architecture Technique', 1)
    
    doc.add_paragraph()
    
    # Stack technologique
    tech_stack = doc.add_heading('Stack Technologique', 2)
    tech_text = doc.add_paragraph(
        'L\'application est développée en utilisant des technologies modernes et robustes :'
    )
    
    tech_list = [
        'Langage de programmation : Python 3.8+',
        'Interface graphique : Tkinter (via ttk)',
        'Gestion des données : JSON',
        'Export de données : CSV',
        'Dépendances : Aucune dépendance externe requise'
    ]
    
    for tech in tech_list:
        p = doc.add_paragraph(tech)
        p.style = 'List Bullet'
    
    doc.add_paragraph()
    
    # Architecture du projet
    architecture = doc.add_heading('Architecture du Projet', 2)
    arch_text = doc.add_paragraph(
        'Le projet suit une architecture modulaire bien structurée :'
    )
    
    arch_structure = [
        'main.py : Point d\'entrée de l\'application',
        'models/ : Modèles de données (Student, Building, Group)',
        'services/ : Services métier (DataManager, Exporter, Scheduler)',
        'interface/ : Interface utilisateur (GUI)',
        'data/ : Stockage des données (JSON, CSV)',
        'constants.py : Constantes et configurations'
    ]
    
    for arch in arch_structure:
        p = doc.add_paragraph(arch)
        p.style = 'List Bullet'

def create_main_features_section(doc):
    """Créer la section fonctionnalités principales"""
    
    doc.add_page_break()
    title = doc.add_heading('4. Fonctionnalités Principales', 1)
    
    doc.add_paragraph()
    
    # Gestion des rôles
    roles = doc.add_heading('4.1 Gestion des Rôles et Authentification', 2)
    roles_text = doc.add_paragraph(
        'L\'application propose un système de rôles hiérarchique :'
    )
    
    roles_list = [
        'Administrateur : Contrôle total sur la gestion des bâtiments, utilisateurs et paramètres globaux',
        'Gestionnaire de bâtiment : Gestion complète de son bâtiment assigné, création de groupes, suivi des tâches',
        'Étudiant/Invité : Accès en lecture seule aux plannings, classements et notifications'
    ]
    
    for role in roles_list:
        p = doc.add_paragraph(role)
        p.style = 'List Bullet'
    
    doc.add_paragraph()
    
    # Gestion intelligente
    smart_management = doc.add_heading('4.2 Gestion Intelligente', 2)
    smart_text = doc.add_paragraph(
        'Fonctionnalités avancées de gestion automatique :'
    )
    
    smart_list = [
        'Planification automatique : Génération de plannings hebdomadaires avec rotation intelligente des tâches',
        'Groupes automatiques : Création automatique de groupes de nettoyage basés sur les blocs et étudiants disponibles',
        'Suivi en temps réel : Tableau de bord interactif pour surveiller la progression, le statut et la qualité du travail'
    ]
    
    for smart in smart_list:
        p = doc.add_paragraph(smart)
        p.style = 'List Bullet'
    
    doc.add_paragraph()
    
    # Gamification
    gamification = doc.add_heading('4.3 Gamification et Motivation', 2)
    gamification_text = doc.add_paragraph(
        'Système de motivation et d\'engagement :'
    )
    
    gamification_list = [
        'Système de badges : Attribution de badges pour encourager la constance, la ponctualité et la qualité',
        'Classements de performance : Affichage des tableaux des meilleurs groupes et étudiants',
        'Notifications : Système de notifications pour informer des tâches assignées, validations et récompenses'
    ]
    
    for gam in gamification_list:
        p = doc.add_paragraph(gam)
        p.style = 'List Bullet'
    
    doc.add_paragraph()
    
    # Rapports et exports
    reports = doc.add_heading('4.4 Rapports et Exports', 2)
    reports_text = doc.add_paragraph(
        'Fonctionnalités de reporting et d\'export :'
    )
    
    reports_list = [
        'Exports CSV : Export facile des plannings, listes d\'étudiants et rapports de performance',
        'Sauvegarde et restauration : Outils intégrés pour sauvegarder et restaurer les données de l\'application'
    ]
    
    for report in reports_list:
        p = doc.add_paragraph(report)
        p.style = 'List Bullet'

def create_user_interface_section(doc):
    """Créer la section interface utilisateur"""
    
    doc.add_page_break()
    title = doc.add_heading('5. Interface Utilisateur', 1)
    
    doc.add_paragraph()
    
    # Design moderne
    design = doc.add_heading('5.1 Design Moderne', 2)
    design_text = doc.add_paragraph(
        'L\'interface utilisateur a été conçue avec un design moderne et professionnel :'
    )
    
    design_list = [
        'Interface graphique basée sur Tkinter avec le thème moderne "clam"',
        'Disposition intuitive avec navigation par onglets',
        'Couleurs professionnelles et cohérentes',
        'Responsive design adapté aux différentes tailles d\'écran',
        'Icônes et éléments visuels pour améliorer l\'expérience utilisateur'
    ]
    
    for design_item in design_list:
        p = doc.add_paragraph(design_item)
        p.style = 'List Bullet'
    
    doc.add_paragraph()
    
    # Navigation
    navigation = doc.add_heading('5.2 Navigation et Ergonomie', 2)
    navigation_text = doc.add_paragraph(
        'L\'ergonomie a été optimisée pour une utilisation efficace :'
    )
    
    navigation_list = [
        'Menu de navigation principal avec accès rapide aux fonctionnalités',
        'Tableau de bord centralisé avec vue d\'ensemble des activités',
        'Formulaires intuitifs pour la saisie et la modification des données',
        'Messages d\'erreur et de confirmation clairs',
        'Raccourcis clavier pour les actions fréquentes'
    ]
    
    for nav in navigation_list:
        p = doc.add_paragraph(nav)
        p.style = 'List Bullet'

def create_data_models_section(doc):
    """Créer la section modèles de données"""
    
    doc.add_page_break()
    title = doc.add_heading('6. Modèles de Données', 1)
    
    doc.add_paragraph()
    
    # Modèle Student
    student_model = doc.add_heading('6.1 Modèle Student', 2)
    student_text = doc.add_paragraph(
        'Le modèle Student représente un étudiant avec ses informations personnelles et données de nettoyage :'
    )
    
    student_attributes = [
        'id : Identifiant unique de l\'étudiant',
        'name : Nom complet de l\'étudiant',
        'building_id : ID du bâtiment de résidence',
        'block : Bloc de résidence (A, B, etc.)',
        'room_number : Numéro de chambre',
        'phone : Numéro de téléphone (optionnel)',
        'email : Adresse email (optionnel)',
        'assigned_groups : Groupes de nettoyage assignés',
        'badges : Badges obtenus pour les performances',
        'last_activity : Dernière activité enregistrée'
    ]
    
    for attr in student_attributes:
        p = doc.add_paragraph(attr)
        p.style = 'List Bullet'
    
    doc.add_paragraph()
    
    # Modèle Building
    building_model = doc.add_heading('6.2 Modèle Building', 2)
    building_text = doc.add_paragraph(
        'Le modèle Building représente un bâtiment de résidence avec sa structure et gestion :'
    )
    
    building_attributes = [
        'id : Identifiant unique du bâtiment',
        'name : Nom du bâtiment',
        'chief_id : ID du chef de bâtiment',
        'blocks : Liste des blocs (A, B, etc.)',
        'rooms_per_block : Nombre de chambres par bloc',
        'people_per_room : Nombre de personnes par chambre',
        'students : Liste des étudiants résidents',
        'room_assignments : Assignations chambre-étudiants',
        'cleaning_groups : Groupes de nettoyage du bâtiment',
        'current_schedule : Planning actuel de nettoyage',
        'overall_completion_rate : Taux de completion global'
    ]
    
    for attr in building_attributes:
        p = doc.add_paragraph(attr)
        p.style = 'List Bullet'
    
    doc.add_paragraph()
    
    # Modèle CleaningGroup
    group_model = doc.add_heading('6.3 Modèle CleaningGroup', 2)
    group_text = doc.add_paragraph(
        'Le modèle CleaningGroup représente un groupe de nettoyage :'
    )
    
    group_attributes = [
        'id : Identifiant unique du groupe',
        'building_id : ID du bâtiment associé',
        'members : Liste des membres du groupe',
        'assigned_areas : Zones de nettoyage assignées',
        'schedule : Planning du groupe',
        'completed_tasks : Tâches complétées',
        'performance_score : Score de performance du groupe'
    ]
    
    for attr in group_attributes:
        p = doc.add_paragraph(attr)
        p.style = 'List Bullet'

def create_services_section(doc):
    """Créer la section services et gestion des données"""
    
    doc.add_page_break()
    title = doc.add_heading('7. Services et Gestion des Données', 1)
    
    doc.add_paragraph()
    
    # DataManager
    data_manager = doc.add_heading('7.1 DataManager', 2)
    data_manager_text = doc.add_paragraph(
        'Service principal pour la gestion de la persistance des données :'
    )
    
    data_manager_features = [
        'Chargement et sauvegarde automatique des données JSON',
        'Gestion des utilisateurs et authentification',
        'CRUD complet pour les étudiants, bâtiments et groupes',
        'Système de notifications et badges',
        'Sauvegarde et restauration des données',
        'Validation et intégrité des données'
    ]
    
    for feature in data_manager_features:
        p = doc.add_paragraph(feature)
        p.style = 'List Bullet'
    
    doc.add_paragraph()
    
    # Exporter
    exporter = doc.add_heading('7.2 Service Exporter', 2)
    exporter_text = doc.add_paragraph(
        'Service pour l\'export des données :'
    )
    
    exporter_features = [
        'Export des plannings en format CSV',
        'Export des listes d\'étudiants',
        'Génération de rapports de performance',
        'Export des statistiques de nettoyage',
        'Formatage personnalisé des données exportées'
    ]
    
    for feature in exporter_features:
        p = doc.add_paragraph(feature)
        p.style = 'List Bullet'
    
    doc.add_paragraph()
    
    # Scheduler
    scheduler = doc.add_heading('7.3 Service Scheduler', 2)
    scheduler_text = doc.add_paragraph(
        'Service de planification automatique :'
    )
    
    scheduler_features = [
        'Génération automatique des plannings hebdomadaires',
        'Rotation intelligente des tâches',
        'Optimisation de la répartition des charges',
        'Gestion des contraintes et disponibilités',
        'Mise à jour automatique des plannings'
    ]
    
    for feature in scheduler_features:
        p = doc.add_paragraph(feature)
        p.style = 'List Bullet'

def create_gamification_section(doc):
    """Créer la section système de gamification"""
    
    doc.add_page_break()
    title = doc.add_heading('8. Système de Gamification', 1)
    
    doc.add_paragraph()
    
    # Système de badges
    badges = doc.add_heading('8.1 Système de Badges', 2)
    badges_text = doc.add_paragraph(
        'Système de récompenses pour motiver les étudiants :'
    )
    
    badge_types = [
        'Badge de Constance : Pour la participation régulière',
        'Badge de Ponctualité : Pour le respect des horaires',
        'Badge de Qualité : Pour l\'excellence dans le travail',
        'Badge de Leadership : Pour les étudiants qui motivent les autres',
        'Badge d\'Innovation : Pour les suggestions d\'amélioration'
    ]
    
    for badge in badge_types:
        p = doc.add_paragraph(badge)
        p.style = 'List Bullet'
    
    doc.add_paragraph()
    
    # Classements
    rankings = doc.add_heading('8.2 Système de Classements', 2)
    rankings_text = doc.add_paragraph(
        'Classements pour encourager la compétition saine :'
    )
    
    ranking_features = [
        'Classement des meilleurs groupes par bâtiment',
        'Classement individuel des étudiants',
        'Classement par performance mensuelle',
        'Affichage des statistiques de progression',
        'Historique des performances'
    ]
    
    for ranking in ranking_features:
        p = doc.add_paragraph(ranking)
        p.style = 'List Bullet'
    
    doc.add_paragraph()
    
    # Notifications
    notifications = doc.add_heading('8.3 Système de Notifications', 2)
    notifications_text = doc.add_paragraph(
        'Système de communication et d\'information :'
    )
    
    notification_types = [
        'Notifications de tâches assignées',
        'Rappels de plannings',
        'Notifications de validation de tâches',
        'Annonces de nouveaux badges',
        'Notifications de mise à jour des classements',
        'Messages d\'encouragement et de motivation'
    ]
    
    for notif in notification_types:
        p = doc.add_paragraph(notif)
        p.style = 'List Bullet'

def create_installation_section(doc):
    """Créer la section installation et utilisation"""
    
    doc.add_page_break()
    title = doc.add_heading('9. Installation et Utilisation', 1)
    
    doc.add_paragraph()
    
    # Prérequis
    prerequisites = doc.add_heading('9.1 Prérequis', 2)
    prerequisites_text = doc.add_paragraph(
        'Avant d\'installer l\'application, assurez-vous d\'avoir :'
    )
    
    prereq_list = [
        'Python 3.8 ou version supérieure installé',
        'Accès en lecture/écriture au dossier d\'installation',
        'Espace disque suffisant pour les données (minimum 10 MB)'
    ]
    
    for prereq in prereq_list:
        p = doc.add_paragraph(prereq)
        p.style = 'List Bullet'
    
    doc.add_paragraph()
    
    # Installation
    installation = doc.add_heading('9.2 Installation', 2)
    installation_text = doc.add_paragraph(
        'Étapes d\'installation :'
    )
    
    install_steps = [
        'Télécharger ou cloner le projet depuis le repository',
        'Naviguer vers le dossier du projet',
        'Vérifier que Python 3.8+ est installé',
        'Lancer l\'application avec la commande : python main.py'
    ]
    
    for step in install_steps:
        p = doc.add_paragraph(step)
        p.style = 'List Number'
    
    doc.add_paragraph()
    
    # Utilisation
    usage = doc.add_heading('9.3 Utilisation', 2)
    usage_text = doc.add_paragraph(
        'Guide d\'utilisation de base :'
    )
    
    usage_steps = [
        'Se connecter avec les identifiants par défaut (admin/admin123)',
        'Configurer les bâtiments et ajouter les étudiants',
        'Créer les groupes de nettoyage ou utiliser la génération automatique',
        'Générer les plannings de nettoyage',
        'Suivre les performances et attribuer les badges',
        'Exporter les rapports selon les besoins'
    ]
    
    for step in usage_steps:
        p = doc.add_paragraph(step)
        p.style = 'List Number'

def create_team_section(doc):
    """Créer la section équipe de développement"""
    
    doc.add_page_break()
    title = doc.add_heading('10. Équipe de Développement', 1)
    
    doc.add_paragraph()
    
    # Présentation de l'équipe
    team_intro = doc.add_paragraph(
        'Ce projet a été développé par une équipe passionnée d\'étudiants en informatique, '
        'déterminés à moderniser la gestion des résidences universitaires.'
    )
    
    doc.add_paragraph()
    
    # Tableau de l'équipe
    team_table = doc.add_table(rows=1, cols=2)
    team_table.style = 'Table Grid'
    
    # En-têtes
    header_cells = team_table.rows[0].cells
    header_cells[0].text = 'Nom'
    header_cells[1].text = 'Prénom(s)'
    
    # Style des en-têtes
    for cell in header_cells:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(12)
    
    # Membres de l'équipe
    team_members = [
        ('KOALAGA', 'Héliane Elichebat'),
        ('NIKIEMA', 'Mariata'),
        ('ZONGO', 'Abdel Sadek Nerebewendé'),
        ('YANOGO', 'Fabiana'),
        ('ZOUNGRANA', 'Pengdwendé Sébastien')
    ]
    
    for last_name, first_name in team_members:
        row_cells = team_table.add_row().cells
        row_cells[0].text = last_name
        row_cells[1].text = first_name
    
    doc.add_paragraph()
    
    # Rôles et responsabilités
    roles = doc.add_heading('10.1 Rôles et Responsabilités', 2)
    roles_text = doc.add_paragraph(
        'Chaque membre de l\'équipe a contribué de manière significative au projet :'
    )
    
    roles_list = [
        'Conception et architecture du système',
        'Développement des modèles de données',
        'Implémentation de l\'interface utilisateur',
        'Création des services de gestion des données',
        'Mise en place du système de gamification',
        'Tests et validation des fonctionnalités',
        'Documentation et présentation du projet'
    ]
    
    for role in roles_list:
        p = doc.add_paragraph(role)
        p.style = 'List Bullet'

def create_conclusion_section(doc):
    """Créer la section conclusion"""
    
    doc.add_page_break()
    title = doc.add_heading('11. Conclusion', 1)
    
    doc.add_paragraph()
    
    # Résumé des réalisations
    achievements = doc.add_heading('11.1 Résumé des Réalisations', 2)
    achievements_text = doc.add_paragraph(
        'Le projet CleanCampus Manager a permis de développer une solution complète et moderne '
        'pour la gestion du nettoyage des résidences universitaires :'
    )
    
    achievements_list = [
        'Application desktop fonctionnelle et intuitive',
        'Système de gestion multi-rôles sécurisé',
        'Automatisation complète des processus de planification',
        'Système de gamification innovant pour motiver les étudiants',
        'Interface utilisateur moderne et professionnelle',
        'Architecture modulaire et extensible',
        'Gestion robuste des données avec sauvegarde automatique'
    ]
    
    for achievement in achievements_list:
        p = doc.add_paragraph(achievement)
        p.style = 'List Bullet'
    
    doc.add_paragraph()
    
    # Impact et bénéfices
    impact = doc.add_heading('11.2 Impact et Bénéfices', 2)
    impact_text = doc.add_paragraph(
        'L\'implémentation de cette solution apporte de nombreux bénéfices :'
    )
    
    impact_list = [
        'Réduction significative du temps de gestion administrative',
        'Amélioration de la motivation et de l\'engagement des étudiants',
        'Optimisation de la répartition des tâches de nettoyage',
        'Amélioration de la qualité et de la régularité du nettoyage',
        'Facilitation de la communication entre les parties prenantes',
        'Modernisation des processus de gestion des résidences'
    ]
    
    for impact_item in impact_list:
        p = doc.add_paragraph(impact_item)
        p.style = 'List Bullet'
    
    doc.add_paragraph()
    
    # Perspectives d'évolution
    future = doc.add_heading('11.3 Perspectives d\'Évolution', 2)
    future_text = doc.add_paragraph(
        'Le projet offre de nombreuses possibilités d\'évolution et d\'amélioration :'
    )
    
    future_list = [
        'Développement d\'une version web/mobile',
        'Intégration avec d\'autres systèmes de gestion universitaire',
        'Ajout de fonctionnalités de reporting avancées',
        'Implémentation d\'un système de notifications push',
        'Intégration de l\'intelligence artificielle pour l\'optimisation',
        'Extension à d\'autres types de résidences et établissements'
    ]
    
    for future_item in future_list:
        p = doc.add_paragraph(future_item)
        p.style = 'List Bullet'
    
    doc.add_paragraph()
    
    # Message de clôture
    closing = doc.add_paragraph()
    closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
    closing_run = closing.add_run('Merci pour votre attention !')
    closing_run.font.size = Pt(14)
    closing_run.font.bold = True
    closing_run.font.italic = True

if __name__ == "__main__":
    try:
        filename = create_clean_campus_presentation()
        print(f"\n✅ Présentation créée avec succès : {filename}")
        print("📄 Le document Word est prêt pour votre présentation !")
    except Exception as e:
        print(f"❌ Erreur lors de la création de la présentation : {e}")
        print("💡 Assurez-vous d'avoir installé python-docx : pip install python-docx") 